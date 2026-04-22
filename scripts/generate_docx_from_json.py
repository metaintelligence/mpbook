from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Inches, Pt
from latex2mathml.converter import convert as latex_to_mathml
from mathml2omml import convert as mathml_to_omml


FORMULA_PATTERN = re.compile(r"\\\((.+?)\\\)")
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
DEFAULT_CROP_MARGIN_PX = 12

# Per-source visual extraction/layout config for this PoC sample set.
# crop: [left, top, right, bottom] in source-image pixel coordinates.
# margin_px: crop bounds padding (pixels) to avoid clipping edges/labels.
VISUAL_LAYOUTS = {
    "sample_input/sample_blank_1.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.8,
        "right_col_width_in": 2.0,
        "items": [
            {"crop": [300, 118, 467, 264], "width_in": 1.9},
        ],
    },
    "sample_input/sample_blank_2.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.8,
        "right_col_width_in": 2.0,
        "items": [
            {"crop": [332, 30, 472, 193], "width_in": 1.9},
        ],
    },
    "sample_input/sample_blank_3.png": {
        "mode": "inline",
        "margin_px": 12,
        "items": [
            {"crop": [28, 68, 254, 227], "width_in": 2.8, "align": "left", "insert_after_body_index": 0},
            {"crop": [20, 255, 262, 470], "width_in": 2.8, "align": "left", "insert_after_body_index": 0},
        ],
    },
    "sample_input/sample_graph_1.png": {
        "mode": "inline",
        "margin_px": 16,
        "items": [
            {"crop": [118, 185, 320, 350], "width_in": 2.9, "align": "center", "insert_after_body_index": 0},
        ],
    },
    "sample_input/sample_graph_2.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.9,
        "right_col_width_in": 1.9,
        "items": [
            {"crop": [330, 22, 468, 190], "width_in": 1.75},
        ],
    },
    "sample_input/sample_graph_3.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.8,
        "right_col_width_in": 2.0,
        "items": [
            {"crop": [316, 18, 478, 215], "width_in": 1.9},
        ],
    },
    "sample_input/sample_graph_4.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.8,
        "right_col_width_in": 2.0,
        "items": [
            {"crop": [292, 30, 474, 208], "width_in": 1.95},
        ],
    },
    "sample_input/sample_graph_5.png": {
        "mode": "right_image",
        "margin_px": 14,
        "left_col_width_in": 4.8,
        "right_col_width_in": 2.0,
        "items": [
            {"crop": [292, 18, 466, 206], "width_in": 1.95},
        ],
    },
    "sample_input/sample_table_1.png": {
        "mode": "inline",
        "margin_px": 10,
        "items": [
            {"crop": [18, 82, 476, 208], "width_in": 5.8, "align": "center", "insert_after_body_index": 0},
        ],
    },
    "sample_input/sample_table_2.png": {
        "mode": "inline",
        "margin_px": 10,
        "items": [
            {"crop": [16, 56, 462, 200], "width_in": 5.6, "align": "center", "insert_after_body_index": 0},
        ],
    },
}


def _add_title(document: Document, title: str) -> None:
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _to_omml_xml(latex_expr: str) -> str:
    mathml_xml = latex_to_mathml(latex_expr)
    omml_xml = mathml_to_omml(mathml_xml).strip()

    if omml_xml.startswith("<m:oMathPara>"):
        return omml_xml.replace(
            "<m:oMathPara>",
            f'<m:oMathPara xmlns:m="{OMML_NS}">',
            1,
        )
    if omml_xml.startswith("<m:oMath>"):
        return omml_xml.replace(
            "<m:oMath>",
            f'<m:oMath xmlns:m="{OMML_NS}">',
            1,
        )
    return f'<m:oMath xmlns:m="{OMML_NS}">{omml_xml}</m:oMath>'


def _add_text_with_omml(paragraph, text: str) -> None:
    cursor = 0
    for match in FORMULA_PATTERN.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])

        latex_expr = match.group(1).strip()
        try:
            omml_elem = parse_xml(_to_omml_xml(latex_expr))
            paragraph._element.append(omml_elem)
        except Exception:
            paragraph.add_run(f"\\({latex_expr}\\)")

        cursor = match.end()

    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _add_formula_aware_paragraph(parent, text: str, space_after_pt: int = 4):
    paragraph = parent.add_paragraph()
    _add_text_with_omml(paragraph, text)
    paragraph.paragraph_format.space_after = Pt(space_after_pt)
    return paragraph


def _add_formula_lines_to_cell(cell, lines: list[str]) -> None:
    cell.text = ""
    if not lines:
        return

    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        _add_text_with_omml(p, line)
        p.paragraph_format.space_after = Pt(4)


def _add_picture_paragraph(parent, image_path: Path, width_in: float, align: str = "center") -> None:
    p = parent.add_paragraph()
    if align == "right":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align == "left":
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    p.paragraph_format.space_after = Pt(4)


def _normalized_crop(
    box: list[int],
    width: int,
    height: int,
    margin_px: int = 0,
) -> tuple[int, int, int, int]:
    left, top, right, bottom = box
    margin = max(0, int(margin_px))
    left -= margin
    top -= margin
    right += margin
    bottom += margin
    left = max(0, min(left, width - 1))
    top = max(0, min(top, height - 1))
    right = max(left + 1, min(right, width))
    bottom = max(top + 1, min(bottom, height))
    return left, top, right, bottom


def _materialize_visual_items(
    qid: str,
    source_image_key: str,
    source_image_path: Path,
    visual_root: Path,
) -> tuple[dict | None, list[dict]]:
    layout_cfg = VISUAL_LAYOUTS.get(source_image_key)
    if not layout_cfg or not source_image_path.exists():
        return None, []

    visual_root.mkdir(parents=True, exist_ok=True)
    created_items: list[dict] = []

    with Image.open(source_image_path) as image:
        img_w, img_h = image.size
        layout_margin = int(layout_cfg.get("margin_px", DEFAULT_CROP_MARGIN_PX))
        for idx, item in enumerate(layout_cfg.get("items", []), start=1):
            item_margin = int(item.get("margin_px", layout_margin))
            crop_box = _normalized_crop(item["crop"], img_w, img_h, margin_px=item_margin)
            cropped = image.crop(crop_box)
            out_path = visual_root / f"{qid}_{idx}_{source_image_path.stem}.png"
            cropped.save(out_path, format="PNG")

            created_items.append(
                {
                    "path": out_path,
                    "width_in": float(item.get("width_in", 2.0)),
                    "align": item.get("align", "center"),
                    "insert_after_body_index": int(item.get("insert_after_body_index", 0)),
                }
            )

    return layout_cfg, created_items


def _add_body_with_inline_visuals(document: Document, body: list[str], visual_items: list[dict]) -> None:
    after_map: dict[int, list[dict]] = {}
    for item in visual_items:
        after_idx = item["insert_after_body_index"]
        after_map.setdefault(after_idx, []).append(item)

    for idx, line in enumerate(body):
        _add_formula_aware_paragraph(document, line)
        for visual in after_map.get(idx, []):
            _add_picture_paragraph(
                document,
                visual["path"],
                width_in=visual["width_in"],
                align=visual["align"],
            )


def _add_body_with_right_visuals(
    document: Document,
    body: list[str],
    visual_items: list[dict],
    left_col_width_in: float,
    right_col_width_in: float,
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(left_col_width_in)
    table.columns[1].width = Inches(right_col_width_in)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    _add_formula_lines_to_cell(left_cell, body)

    right_cell.text = ""
    if visual_items:
        first = True
        for visual in visual_items:
            if first:
                p = right_cell.paragraphs[0]
                first = False
            else:
                p = right_cell.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run()
            run.add_picture(str(visual["path"]), width=Inches(visual["width_in"]))
            p.paragraph_format.space_after = Pt(4)

    document.add_paragraph("")


def _add_question(
    document: Document,
    question: dict,
    base_dir: Path,
    visual_root: Path,
) -> None:
    qid = question.get("id", "Q?")
    qtype = question.get("type", "unknown")
    title = question.get("title", "")
    source_image_key = question.get("source_image", "")
    source_image_path = base_dir / source_image_key
    body = question.get("body", [])
    choices = question.get("choices", [])

    document.add_heading(f"{qid} [{qtype}] {title}", level=1)

    layout_cfg, visual_items = _materialize_visual_items(
        qid=qid,
        source_image_key=source_image_key,
        source_image_path=source_image_path,
        visual_root=visual_root,
    )

    if layout_cfg and layout_cfg.get("mode") == "right_image" and visual_items:
        _add_body_with_right_visuals(
            document=document,
            body=body,
            visual_items=visual_items,
            left_col_width_in=float(layout_cfg.get("left_col_width_in", 4.8)),
            right_col_width_in=float(layout_cfg.get("right_col_width_in", 2.0)),
        )
    elif layout_cfg and layout_cfg.get("mode") == "inline" and visual_items:
        _add_body_with_inline_visuals(document, body, visual_items)
    else:
        for line in body:
            _add_formula_aware_paragraph(document, line)

    if choices:
        document.add_paragraph("선지:")
        for idx, choice in enumerate(choices, start=1):
            p = document.add_paragraph()
            p.add_run(f"{idx}. ")
            _add_text_with_omml(p, choice)
            p.paragraph_format.space_after = Pt(2)


def build_docx(input_json_path: Path, output_docx_path: Path) -> None:
    payload = json.loads(input_json_path.read_text(encoding="utf-8"))
    questions = payload.get("questions", [])

    document = Document()
    _add_title(document, "Math Problem Book (PoC)")
    document.add_paragraph("This document was generated from structured JSON.")
    document.add_paragraph(
        "Inline formulas wrapped by \\(...\\) are converted to Word equation objects (OMML)."
    )
    document.add_paragraph(
        "Tables/diagrams are cropped from source images and reinserted at question layout positions."
    )
    document.add_paragraph("")

    base_dir = input_json_path.parent.parent
    visual_root = output_docx_path.parent / "derived_visuals"

    for question in questions:
        _add_question(document, question, base_dir, visual_root)
        document.add_page_break()

    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx_path))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate a worksheet DOCX from structured math-question JSON with OMML formulas and visual crops."
    )
    parser.add_argument(
        "--input",
        default="output/structured_questions.json",
        help="Path to input JSON file.",
    )
    parser.add_argument(
        "--output",
        default="output/math_problem_book.docx",
        help="Path to output DOCX file.",
    )
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    build_docx(input_path, output_path)
    print(f"Generated DOCX: {output_path}")


if __name__ == "__main__":
    main()
